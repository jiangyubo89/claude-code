import re, os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Read the markdown file
with open('N-MP/manuscript/Chapter8_Identification_and_Tracking_Tools.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def set_run_font(run, name='Times New Roman', size=Pt(12), bold=False, italic=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def parse_rich_text(p, text):
    """Parse markdown bold/italic and add runs to paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            set_run_font(run, italic=True)
        else:
            run = p.add_run(part)
            set_run_font(run)

def add_table_from_md(doc, header_line, rows_lines):
    """Create a Word table from markdown table lines."""
    headers = [c.strip() for c in header_line.split('|') if c.strip()]
    num_cols = len(headers)

    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        if i < num_cols:
            cell = hdr.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_line in rows_lines:
        cols = [c.strip() for c in row_line.split('|') if c.strip()]
        if len(cols) == 0:
            continue
        row = table.add_row()
        for i, val in enumerate(cols):
            if i < num_cols:
                cell = row.cells[i]
                cell.text = ''
                p = cell.paragraphs[0]
                clean = val.strip('*')
                run = p.add_run(clean)
                if val.startswith('**'):
                    run.font.bold = True
                run.font.size = Pt(8)
                run.font.name = 'Times New Roman'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.autofit = True
    return table

# Parse and build document
lines = md_content.split('\n')
i = 0
in_table = False
table_header = None
table_rows = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip horizontal rules and empty lines
    if stripped == '---' or stripped == '':
        if in_table and table_header:
            add_table_from_md(doc, table_header, table_rows)
            doc.add_paragraph()
            in_table = False
            table_header = None
            table_rows = []
        i += 1
        continue

    # Table detection
    if '|' in stripped and not stripped.startswith('>') and not stripped.startswith('#'):
        if not in_table:
            if i + 1 < len(lines) and re.match(r'^\|[\s:|-]+\|$', lines[i+1].strip()):
                in_table = True
                table_header = stripped
                table_rows = []
                i += 2
                continue
        if in_table:
            table_rows.append(stripped)
            i += 1
            continue
    else:
        if in_table and table_header:
            add_table_from_md(doc, table_header, table_rows)
            doc.add_paragraph()
            in_table = False
            table_header = None
            table_rows = []

    # Headings
    if stripped.startswith('# ') and not stripped.startswith('## '):
        text = stripped[2:].strip()
        add_heading_styled(doc, text, 1)
        i += 1
        continue
    if stripped.startswith('## '):
        text = stripped[3:].strip()
        add_heading_styled(doc, text, 2)
        i += 1
        continue
    if stripped.startswith('### '):
        text = stripped[4:].strip()
        if text.startswith('Table 8.'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, bold=True, size=Pt(10))
        else:
            add_heading_styled(doc, text, 3)
        i += 1
        continue

    # Block quotes
    if stripped.startswith('> '):
        text = stripped[2:].strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        set_run_font(run, size=Pt(10), italic=True)
        i += 1
        continue

    # Bullet list items
    if stripped.startswith('- '):
        text = stripped[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        parse_rich_text(p, text)
        i += 1
        continue

    # Numbered list
    if re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        p = doc.add_paragraph(style='List Number')
        parse_rich_text(p, text)
        i += 1
        continue

    # Math equations
    if stripped.startswith('$$'):
        eq_lines = [stripped.replace('$$', '')]
        i += 1
        while i < len(lines) and '$$' not in lines[i]:
            eq_lines.append(lines[i].strip())
            i += 1
        if i < len(lines):
            eq_lines.append(lines[i].strip().replace('$$', ''))
            i += 1
        eq_text = ' '.join(l for l in eq_lines if l).strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(eq_text)
        set_run_font(run, italic=True, size=Pt(11))
        continue

    # Italic-only lines (footnotes etc)
    if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.strip('*'))
        set_run_font(run, italic=True, size=Pt(10))
        i += 1
        continue

    # Normal paragraph with rich text
    p = doc.add_paragraph()
    parse_rich_text(p, stripped)
    i += 1

# Finish remaining table
if in_table and table_header:
    add_table_from_md(doc, table_header, table_rows)

# Save Word document
output_dir = 'N-MP/Chapter8_Output'
os.makedirs(output_dir, exist_ok=True)
docx_path = os.path.join(output_dir, 'Chapter8_Identification_and_Tracking_Tools.docx')
doc.save(docx_path)
print(f'Word document saved: {docx_path}')

# Copy figures
fig_src = 'meta_analysis_output/optimized_v3'
fig_dir = os.path.join(output_dir, 'Figures')
os.makedirs(fig_dir, exist_ok=True)

for fig in os.listdir(fig_src):
    if fig.endswith('.png'):
        shutil.copy2(os.path.join(fig_src, fig), os.path.join(fig_dir, fig))
        print(f'Copied figure: {fig}')

# Copy manuscript files
shutil.copy2('N-MP/manuscript/Chapter8_Identification_and_Tracking_Tools.md',
             os.path.join(output_dir, 'Chapter8_Identification_and_Tracking_Tools.md'))
shutil.copy2('N-MP/manuscript/Chapter8_References.ris',
             os.path.join(output_dir, 'Chapter8_References.ris'))
print('Copied: manuscript .md and .ris files')

# Copy meta-analysis data
data_dir = os.path.join(output_dir, 'Data')
os.makedirs(data_dir, exist_ok=True)
for f in ['pooled_results.csv', 'meta_analysis_input.csv', 'META_ANALYSIS_REPORT.txt']:
    src = os.path.join(fig_src, f)
    if os.path.exists(src):
        shutil.copy2(src, os.path.join(data_dir, f))
        print(f'Copied data: {f}')

print(f'\nAll files organized in: {output_dir}/')
print('Contents:')
for root, dirs, files in os.walk(output_dir):
    level = root.replace(output_dir, '').count(os.sep)
    indent = '  ' * level
    print(f'{indent}{os.path.basename(root)}/')
    subindent = '  ' * (level + 1)
    for file in files:
        size = os.path.getsize(os.path.join(root, file))
        print(f'{subindent}{file} ({size/1024:.0f} KB)')
