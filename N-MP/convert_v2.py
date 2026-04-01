"""
Robust Markdown-to-Word converter for Chapter 8.
Handles full content, embeds figures, creates complete docx.
"""
import re, os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD_PATH = 'N-MP/manuscript/Chapter8_Identification_and_Tracking_Tools.md'
OUT_DIR = 'N-MP/Chapter8_Output'
FIG_OUT = os.path.join(OUT_DIR, 'Figures')
DOCX_OUT = os.path.join(OUT_DIR, 'Chapter8_Final_Draft.docx')

# Chapter figures: inserted AFTER the heading of the matching section
SECTION_FIGURES = {
    '8.2 Molecular Fingerprinting for iMNP Detection': (
        os.path.join(FIG_OUT, 'Fig8_1_method_comparison.png'),
        'Figure 8.1  Comparative Performance of Analytical Methods for iMNP Detection in Biological Tissues. '
        'Heatmap showing performance scores (0-10) across eight dimensions for eight analytical techniques. '
        'Py-GC/MS and AFM-IR excel at nanoplastics detection; LDIR offers highest throughput; '
        'mu-Raman provides the best spatial resolution among non-destructive methods.'
    ),
    '8.4 Biodistribution Studies': (
        os.path.join(FIG_OUT, 'Fig8_2_biodistribution.png'),
        'Figure 8.2  iMNP Biodistribution in Animal Models (Oral PS nanoparticle exposure, rodent models, 7-28 days). '
        '(A) Organ accumulation hierarchy showing percentage of administered dose retained in each organ. '
        'GI tract retains the majority (>90% not systemically absorbed); liver is the primary site of systemic accumulation. '
        '(B) Size-dependent distribution across three key organs, demonstrating that smaller particles achieve '
        'greater systemic translocation while larger particles are predominantly retained in the GI tract.'
    ),
    '8.5 Biological Barrier Crossing': (
        os.path.join(FIG_OUT, 'Fig8_3_barrier_thresholds.png'),
        'Figure 8.3  Biological Barrier Size Thresholds and Translocation Mechanisms for iMNPs. '
        'Horizontal bars represent the maximum particle size permitting translocation across each barrier. '
        'Estimated translocation efficiency (%) for ~100 nm particles is shown. '
        'Key transport mechanisms for each barrier are annotated.'
    ),
    '8.8 Kinetic Modeling': (
        os.path.join(FIG_OUT, 'Fig8_4_PBPK_model.png'),
        'Figure 8.4  Physiologically Based Pharmacokinetic (PBPK) Model for iMNP Biodistribution. '
        'Compartmental structure showing organ compartments connected by blood flows. '
        'Exposure routes (oral, inhalation, dermal) and excretion pathways (feces >90%, bile, urine) are indicated. '
        'Percentages represent typical organ accumulation from animal studies. '
        'Key model parameters are listed (Qi, Ci, Pi, kclear, Vmax/Km).'
    ),
    '8.9 Human Evidence': (
        os.path.join(FIG_OUT, 'Fig8_5_human_evidence.png'),
        'Figure 8.5  Human Evidence: Detection Prevalence and Study Quality Assessment. '
        '(A) Study-level prevalence estimates with 95% CIs (square size proportional to sample size). '
        'Red diamond represents the overall pooled estimate of 68.5% (95% CI: 56.3-78.6%). '
        '(B) Quality assessment summary showing number of studies (of 11) meeting each Modified NOS criterion. '
        'Red bars highlight criteria met by 2 or fewer studies.'
    ),
}

# No more blockquote figure insertions (meta-analysis plots removed)
BLOCKQUOTE_FIGURES = {}

with open(MD_PATH, 'r', encoding='utf-8') as f:
    md = f.read()

doc = Document()

# ── Page setup ──
for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

# ── Base style ──
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(12)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(4)

def sf(run, sz=Pt(12), bold=False, italic=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = sz
    run.font.bold = bold
    run.font.italic = italic

def add_rich(p, text):
    """Add text with **bold** and *italic* parsing."""
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            sf(r, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            sf(r, italic=True)
        else:
            r = p.add_run(part)
            sf(r)

def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Insert a figure with caption."""
    if img_path and os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        sf(r, sz=Pt(10), italic=True)
        doc.add_paragraph()  # spacing

def make_table(doc, header_line, data_lines):
    """Build a Word table from markdown pipe-delimited lines."""
    headers = [c.strip() for c in header_line.split('|') if c.strip()]
    nc = len(headers)
    tbl = doc.add_table(rows=1, cols=nc)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        pr = cell.paragraphs[0]
        r = pr.add_run(h)
        sf(r, sz=Pt(8), bold=True)
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Shade header
        tc = cell._element.find(qn('w:tcPr'))
        if tc is None:
            tc = cell._element.makeelement(qn('w:tcPr'), {})
            cell._element.insert(0, tc)
        shd = tc.makeelement(qn('w:shd'), {qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'})
        tc.append(shd)

    for dl in data_lines:
        cols = [c.strip() for c in dl.split('|') if c.strip()]
        if not cols:
            continue
        row = tbl.add_row()
        for j, val in enumerate(cols):
            if j >= nc:
                break
            cell = row.cells[j]
            cell.text = ''
            pr = cell.paragraphs[0]
            clean = val.strip('*')
            r = pr.add_run(clean)
            sf(r, sz=Pt(8), bold=val.startswith('**'))
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.autofit = True

# ── Split into lines and process ──
lines = md.split('\n')
idx = 0
N = len(lines)
in_references = False  # Track References section for plain-text numbering

while idx < N:
    line = lines[idx]
    s = line.strip()

    # Skip blank / hr
    if s == '' or s == '---':
        idx += 1
        continue

    # ── Headings ──
    if s.startswith('# ') and not s.startswith('## '):
        h = doc.add_heading(s[2:], level=1)
        for r in h.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0,0,0)
        idx += 1; continue

    if s.startswith('## '):
        sec_title = s[3:]
        h = doc.add_heading(sec_title, level=2)
        for r in h.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0,0,0)
        # Track if we entered References or Figure Legends section
        if 'References' in sec_title or 'Figure Legends' in sec_title:
            in_references = True
        else:
            in_references = False
        # Insert chapter figure after matching section heading
        for sec_key, (fig_path, fig_cap) in SECTION_FIGURES.items():
            if sec_key in sec_title:
                doc.add_paragraph()
                add_figure(doc, fig_path, fig_cap, width=Inches(5.8))
                break
        idx += 1; continue

    if s.startswith('### '):
        txt = s[4:]
        if txt.startswith('Table 8.'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt)
            sf(r, sz=Pt(10), bold=True)
        else:
            h = doc.add_heading(txt, level=3)
            for r in h.runs:
                r.font.name = 'Times New Roman'
                r.font.color.rgb = RGBColor(0,0,0)
        idx += 1; continue

    # ── Tables ──
    if '|' in s and not s.startswith('>') and not s.startswith('#'):
        if idx + 1 < N and re.match(r'^\|[\s:|-]+\|', lines[idx+1].strip()):
            header = s
            idx += 2  # skip header + separator
            data = []
            while idx < N:
                sl = lines[idx].strip()
                if '|' in sl and not sl.startswith('>') and not sl.startswith('#') and sl != '---' and sl != '':
                    data.append(sl)
                    idx += 1
                else:
                    break
            make_table(doc, header, data)
            doc.add_paragraph()
            continue
        # Single pipe line that is not table header? treat as normal
        # fall through

    # ── Block quotes / figure placeholders ──
    if s.startswith('> '):
        text = s[2:].strip()
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)

        # Check if this references a meta-analysis figure we can embed
        inserted = False
        for fig_key, (fig_path, fig_cap) in BLOCKQUOTE_FIGURES.items():
            if fig_key in text and fig_path and os.path.exists(fig_path):
                add_figure(doc, fig_path, fig_cap, width=Inches(5.5))
                inserted = True
                break

        if not inserted:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(clean)
            sf(r, sz=Pt(10), italic=True)
        idx += 1; continue

    # ── Math block ──
    if s.startswith('$$'):
        # Check if equation is all on one line: $$...$$
        if s.endswith('$$') and len(s) > 4:
            eq_text = s[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('[Equation] ' + eq_text)
            sf(r, sz=Pt(11), italic=True)
            idx += 1
        else:
            # Multi-line equation
            eq_parts = []
            if s != '$$':
                eq_parts.append(s.replace('$$', ''))
            idx += 1
            while idx < N:
                sl = lines[idx].strip()
                if '$$' in sl:
                    eq_parts.append(sl.replace('$$', ''))
                    idx += 1
                    break
                eq_parts.append(sl)
                idx += 1
            eq_text = ' '.join(p2 for p2 in eq_parts if p2).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('[Equation] ' + eq_text)
            sf(r, sz=Pt(11), italic=True)

        # Also consume the "where ..." line if it follows
        if idx < N and lines[idx].strip().startswith('where '):
            wp = doc.add_paragraph()
            add_rich(wp, lines[idx].strip())
            idx += 1
        continue

    # ── Bullet list ──
    if s.startswith('- '):
        text = s[2:]
        p = doc.add_paragraph(style='List Bullet')
        add_rich(p, text)
        idx += 1; continue

    # ── Numbered list ──
    m = re.match(r'^(\d+)\.\s+(.*)', s)
    if m:
        num = m.group(1)
        text = m.group(2)
        if in_references:
            # References: use plain paragraph with number preserved, smaller font
            p = doc.add_paragraph()
            r = p.add_run(f'{num}. {text}')
            sf(r, sz=Pt(10))
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
        else:
            # Normal numbered list
            p = doc.add_paragraph(style='List Number')
            add_rich(p, text)
        idx += 1; continue

    # ── Italic-only footnote lines ──
    if s.startswith('*') and s.endswith('*') and not s.startswith('**'):
        p = doc.add_paragraph()
        r = p.add_run(s.strip('*'))
        sf(r, sz=Pt(10), italic=True)
        idx += 1; continue

    # ── Normal paragraph ──
    p = doc.add_paragraph()
    add_rich(p, s)
    idx += 1

# ── Save ──
os.makedirs(OUT_DIR, exist_ok=True)
doc.save(DOCX_OUT)

# Verify
ndoc = Document(DOCX_OUT)
print(f'Saved: {DOCX_OUT}')
print(f'Total paragraphs: {len(ndoc.paragraphs)}')

# Check key sections present
full_text = '\n'.join(p.text for p in ndoc.paragraphs)
checks = ['8.9 Human Evidence', '8.10 Comparative Analysis', '8.11 Conclusions',
           'References', 'Figure Legends', '68.5%']
for c in checks:
    if c in full_text:
        print(f'  OK: "{c}" found')
    else:
        print(f'  MISSING: "{c}"')

# Count images
img_count = 0
for rel in ndoc.part.rels.values():
    if 'image' in rel.reltype:
        img_count += 1
print(f'Embedded images: {img_count}')

fsize = os.path.getsize(DOCX_OUT)
print(f'File size: {fsize/1024:.0f} KB')
